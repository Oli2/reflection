from docx import Document
import PyPDF2
import os
import io

def read_document(file_data) -> str:
    """Read content from either DOCX or PDF files."""
    try:
        # Handle bytes data from Gradio
        if isinstance(file_data, bytes):
            file_obj = io.BytesIO(file_data)
            
            # Try to read as PDF first
            try:
                pdf_reader = PyPDF2.PdfReader(file_obj)
                return '\n'.join(page.extract_text() for page in pdf_reader.pages)
            except:
                # If PDF fails, try DOCX
                file_obj.seek(0)
                doc = Document(file_obj)
                return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
                
        # Handle file object (for backward compatibility)
        elif hasattr(file_data, 'read'):
            return read_pdf(file_data) if is_pdf(file_data) else read_docx(file_data)
        else:
            raise ValueError("Unsupported file format or type")
            
    except Exception as e:
        raise ValueError("Unable to read document. Please ensure it's a valid PDF or DOCX file.")

def is_pdf(file_obj):
    """Check if file is PDF by reading first few bytes"""
    try:
        file_obj.seek(0)
        header = file_obj.read(4)
        file_obj.seek(0)
        return header.startswith(b'%PDF')
    except:
        return False

def read_pdf(file) -> str:
    """Read content from PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        return '\n'.join(page.extract_text() for page in pdf_reader.pages)
    except Exception as e:
        raise ValueError("Unable to read PDF file. Please ensure it's a valid PDF document.")

def read_docx(file) -> str:
    """Read content from DOCX file."""
    try:
        doc = Document(file)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    except Exception as e:
        raise ValueError("Unable to read DOCX file. Please ensure it's a valid Word document.") 